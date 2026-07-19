import io
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches

def export_txt(content: str) -> bytes:
    """Returns raw text as bytes."""
    return content.encode("utf-8")

def export_pdf(title: str, data: any, format_type: str) -> bytes:
    """
    Exports the given data to a professional PDF layout.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    pdf.set_font("Helvetica", style="B", size=16)
    pdf.cell(0, 10, title, ln=True, align="C")
    pdf.ln(5)
    
    pdf.set_font("Helvetica", size=11)
    
    import textwrap
    import re
    # Helper to clean text
    def safe_text(txt):
        if not txt:
            return ""
        # Remove markdown bold/italic asterisks
        s = re.sub(r'[*_]', '', str(txt))
        # Replace common unicode chars
        s = s.replace("\t", "    ").replace("—", "-").replace("‘", "'").replace("’", "'").replace("“", '"').replace("”", '"')
        s = s.replace("\xa0", " ") # replace non-breaking spaces
        s = s.encode('latin-1', 'replace').decode('latin-1')
        
        # Wrap at 90 characters to avoid FPDF width calculation issues
        wrapped = textwrap.fill(s, width=90)
        return wrapped

    if format_type == "notes":
        for line in data.split('\n'):
            line = safe_text(line)
            if line.startswith("### "):
                pdf.set_font("Helvetica", style="B", size=12)
                pdf.multi_cell(0, 8, line.replace("### ", ""))
                pdf.set_font("Helvetica", size=11)
            elif line.startswith("## "):
                pdf.set_font("Helvetica", style="B", size=14)
                pdf.multi_cell(0, 10, line.replace("## ", ""))
                pdf.set_font("Helvetica", size=11)
            elif line.startswith("# "):
                pdf.set_font("Helvetica", style="B", size=16)
                pdf.multi_cell(0, 12, line.replace("# ", ""))
                pdf.set_font("Helvetica", size=11)
            else:
                pdf.multi_cell(0, 6, line)
                
    elif format_type == "quiz":
        for i, q in enumerate(data):
            pdf.set_font("Helvetica", style="B", size=11)
            pdf.multi_cell(0, 6, safe_text(f"Q{i+1}: {q.get('question', '')}"))
            pdf.set_font("Helvetica", size=11)
            for opt in q.get('options', []):
                pdf.multi_cell(0, 6, safe_text(f"   {opt}"))
            pdf.ln(2)
            pdf.set_font("Helvetica", style="I", size=10)
            pdf.multi_cell(0, 6, safe_text(f"Answer: {q.get('answer', '')}"))
            pdf.multi_cell(0, 6, safe_text(f"Explanation: {q.get('explanation', '')}"))
            pdf.set_font("Helvetica", size=11)
            pdf.ln(5)

    elif format_type == "flashcards":
        for i, card in enumerate(data):
            pdf.set_font("Helvetica", style="B", size=12)
            pdf.multi_cell(0, 8, safe_text(f"Card {i+1}"))
            pdf.set_font("Helvetica", style="B", size=11)
            pdf.multi_cell(0, 6, safe_text(f"Q: {card.get('question', '')}"))
            pdf.set_font("Helvetica", size=11)
            pdf.multi_cell(0, 6, safe_text(f"A: {card.get('answer', '')}"))
            pdf.ln(5)

    elif format_type == "interview":
        for difficulty in ["Beginner", "Intermediate", "Advanced"]:
            questions = data.get(difficulty, [])
            if questions:
                pdf.set_font("Helvetica", style="B", size=14)
                pdf.multi_cell(0, 10, safe_text(f"{difficulty} Level"))
                pdf.set_font("Helvetica", size=11)
                for i, q in enumerate(questions):
                    pdf.set_font("Helvetica", style="B", size=11)
                    pdf.multi_cell(0, 6, safe_text(f"Q{i+1}: {q.get('question', '')}"))
                    pdf.set_font("Helvetica", size=11)
                    pdf.multi_cell(0, 6, safe_text(f"Suggested Answer: {q.get('answer', '')}"))
                    pdf.ln(4)

    elif format_type == "practice":
        for difficulty in ["Easy", "Medium", "Hard"]:
            questions = data.get(difficulty, [])
            if questions:
                pdf.set_font("Helvetica", style="B", size=14)
                pdf.multi_cell(0, 10, safe_text(f"{difficulty} Level"))
                pdf.set_font("Helvetica", size=11)
                for i, q in enumerate(questions):
                    pdf.set_font("Helvetica", style="B", size=11)
                    pdf.multi_cell(0, 6, safe_text(f"Q{i+1}: {q.get('question', '')}"))
                    pdf.set_font("Helvetica", size=11)
                    pdf.multi_cell(0, 6, safe_text(f"Solution: {q.get('solution', '')}"))
                    pdf.ln(4)

    return pdf.output(dest="S")

def export_docx(title: str, data: any, format_type: str) -> bytes:
    """
    Exports the given data to a structured DOCX.
    """
    doc = Document()
    doc.add_heading(title, 0)
    
    if format_type == "notes":
        for line in data.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith("### "):
                doc.add_heading(line.replace("### ", ""), level=3)
            elif line.startswith("## "):
                doc.add_heading(line.replace("## ", ""), level=2)
            elif line.startswith("# "):
                doc.add_heading(line.replace("# ", ""), level=1)
            elif line.startswith("- "):
                doc.add_paragraph(line.replace("- ", "", 1), style='List Bullet')
            else:
                doc.add_paragraph(line)
                
    elif format_type == "quiz":
        for i, q in enumerate(data):
            p = doc.add_paragraph()
            p.add_run(f"Q{i+1}: {q.get('question', '')}").bold = True
            for opt in q.get('options', []):
                doc.add_paragraph(opt, style='List Bullet')
            doc.add_paragraph(f"Answer: {q.get('answer', '')}", style='Intense Quote')
            doc.add_paragraph(f"Explanation: {q.get('explanation', '')}")
            doc.add_paragraph() # spacing

    elif format_type == "flashcards":
        for i, card in enumerate(data):
            doc.add_heading(f"Card {i+1}", level=2)
            p_q = doc.add_paragraph()
            p_q.add_run("Q: ").bold = True
            p_q.add_run(card.get('question', ''))
            p_a = doc.add_paragraph()
            p_a.add_run("A: ").bold = True
            p_a.add_run(card.get('answer', ''))

    elif format_type in ["interview", "practice"]:
        for difficulty in ["Beginner", "Intermediate", "Advanced", "Easy", "Medium", "Hard"]:
            questions = data.get(difficulty, [])
            if questions:
                doc.add_heading(f"{difficulty} Level", level=2)
                for i, q in enumerate(questions):
                    p = doc.add_paragraph()
                    p.add_run(f"Q{i+1}: {q.get('question', '')}").bold = True
                    
                    ans_key = 'answer' if format_type == 'interview' else 'solution'
                    doc.add_paragraph(q.get(ans_key, ''))

    # Save to in-memory buffer
    f = io.BytesIO()
    doc.save(f)
    return f.getvalue()
